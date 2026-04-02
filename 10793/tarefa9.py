import pickle
from docx import Document
from fpdf import FPDF
import numpy as np
import matplotlib.pyplot as plt

array1 = []
array2 = []

def reg_arrays():
    global array1, array2
    
    try:
        tamanho = int(input("Quantos valores pretende registar em cada array? "))
        
        array1 = []
        array2 = []
        
        print("\nRegisto do primeiro array")
        for i in range(tamanho):
            while True:
                try:
                    valor = float(input(f"Introduza o {i+1}º valor do array1: "))
                    array1.append(valor)
                    break
                except ValueError:
                    print("Erro! Introduza um número válido.")
        
        print("\nRegisto do seundo array")
        for i in range(tamanho):
            while True:
                try:
                    valor = float(input(f"Introduza o {i+1}º valor do array2: "))
                    array2.append(valor)
                    break
                except ValueError:
                    print("Erro! Introduza um número válido.")
                    
        print("\nArrays registados com sucesso!\n")
        
    except ValueError:
        print("Erro! O tamanho deve ser um número inteiro.\n")

def list_arrays():
    if not array1:
        print("Os arrays estão vazios!\n")
    else:
        print(f"Array 1: {array1}")
        print(f"Array 2: {array2}\n")

def soma_arrays():
    if not array1:
        print("Os arrays estão vazios!\n")
    else:
        resultado = np.array(array1) + np.array(array2)
        print(f"Resultado da soma: {resultado}\n")

def sub_arrays():
    if not array1:
        print("Os arrays estão vazios!\n")
    else:
        resultado = np.array(array1) - np.array(array2)
        print(f"Resultado da subtração: {resultado}\n")

def multi_arrays():
    if not array1:
        print("Os arrays estão vazios!\n")
    else:
        resultado = np.array(array1) * np.array(array2)
        print(f"Resultado da multiplicação: {resultado}\n")

def div_arrays():
    if not array1:
        print("Os arrays estão vazios!\n")
    else:
        try:
            resultado = np.array(array1) / np.array(array2)
            print(f"Resultado da divisão: {resultado}\n")
        except ZeroDivisionError:
            print("Erro! Não é possível dividir por zero.\n")

def stat_array():
    if not array1:
        print("Os arrays estão vazios!\n")
    else:
        escolha = input("Escolha o array (1 ou 2): ")
        
        if escolha == "1":
            dados = np.array(array1)
        elif escolha == "2":
            dados = np.array(array2)
        else:
            print("Opção inválida!\n")
            return
        
        print(f"Média: {np.mean(dados)}")
        print(f"Mediana: {np.median(dados)}")
        print(f"Desvio padrão: {np.std(dados)}")
        print(f"Maior valor: {np.max(dados)}")
        print(f"Menor valor: {np.min(dados)}\n")

def exportar_pickle():
    if not array1:
        print("Não existem dados para guardar.\n")
    else:
        with open("operacoes.pkl", "wb") as ficheiro:
            pickle.dump((array1, array2), ficheiro)
        print("Backup exportado com sucesso!\n")

def importar_pickle():
    global array1, array2
    try:
        with open("operacoes.pkl", "rb") as ficheiro:
            array1, array2 = pickle.load(ficheiro)
        print("Backup importado com sucesso!\n")
    except FileNotFoundError:
        print("Ficheiro não encontrado!\n")

def exportar_word():
    if not array1:
        print("Não existem dados para exportar.\n")
    else:
        try:
            doc = Document()
            titulo = input("Indique o título do documento: ")
            doc.add_heading(titulo, 0)
            
            for i, (array1, array2) in enumerate(zip(array1, array2), start=1):
                doc.add_paragraph(f"{i} - {array1} | {array2}")
            
            doc.save("operacoes.docx")
            print("Ficheiro Word gerado com sucesso!\n")
        except Exception as e:
            print(f"Erro ao gerar Word: {e}\n")

def exportar_pdf():
    if not array1:
        print("Não existem dados para exportar.\n")
    else:
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            titulo = input("Indique o título do documento: ")
            pdf.cell(200, 10, txt=titulo, ln=True, align="C")
            
            for i, (array1, array2) in enumerate(zip(array1, array2), start=1):
                pdf.cell(200, 10, txt=f"{i} - {array1} | {array2}", ln=True)
            
            pdf.output("operacoes.pdf")
            print("Ficheiro PDF gerado com sucesso!\n")
        except Exception as e:
            print(f"Erro ao gerar PDF: {e}\n")

def grafico_barras():
    if not array1:
        print("Os arrays estão vazios!\n")
    else:
        x = np.arange(len(array1))
        largura = 0.35

        plt.figure(figsize=(8,5))
        plt.bar(x - largura/2, array1, largura, label='Array 1', color='green')
        plt.bar(x + largura/2, array2, largura, label='Array 2', color='orange')

        plt.xlabel("Índice")
        plt.ylabel("Valores")
        plt.title("Gráfico de Barras Comparativo")
        plt.xticks(x)
        plt.legend()
        plt.tight_layout()
        plt.show()

while True:
    print("Operações com Arrays Numéricos")
    print("1 - Registar valores nos arrays")
    print("2 - Listar arrays")
    print("3 - Somar arrays")
    print("4 - Subtrair arrays")
    print("5 - Multiplicar arrays")
    print("6 - Dividir arrays")
    print("7 - Estatísticas de um array")
    print("8 - Exportar Backup")
    print("9 - Importar Backup")
    print("10 - Exportar para Word")
    print("11 - Exportar para PDF")
    print("12 - Grafico")
    print("0 - Terminar")
    
    opcao = input("Escolha uma opção: ")
    
    if opcao == "1":
        reg_arrays()
    elif opcao == "2":
        list_arrays()
    elif opcao == "3":
        soma_arrays()
    elif opcao == "4":
        sub_arrays()
    elif opcao == "5":
        multi_arrays()
    elif opcao == "6":
        div_arrays()
    elif opcao == "7":
        stat_array()
    elif opcao == "8":
        exportar_pickle()
    elif opcao == "9":
        importar_pickle()
    elif opcao == "10":
        exportar_word()
    elif opcao == "11":
        exportar_pdf()
    elif opcao == "12":
        grafico_barras()
    elif opcao == "0":
        print("Programa terminado. Obrigado!")
        break
    else:
        print("Opção inválida!\n")