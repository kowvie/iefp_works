import pickle
from docx import Document
from fpdf import FPDF
import numpy as np

nomes = []
notas = []

def adicionar_aluno():
    try:
        nome = input("Introduza o nome do aluno: ")
        
        if nome == "":
            print("O nome não pode estar vazio!\n")
            return
        
        nota = float(input("Introduza a nota do aluno (0-20): "))
        
        if 0 <= nota <= 20:
            nomes.append(nome)
            notas.append(nota)
            print("Aluno registado com sucesso!\n")
        else:
            print("Erro! A nota deve estar entre 0 e 20.\n")
            
    except ValueError:
        print("Erro! A nota deve ser um número válido.\n")

def listar_alunos():
    if not nomes:
        print("Não existem alunos registados.\n")
    else:
        print("Lista de Alunos:")
        for i, (nome, nota) in enumerate(zip(nomes, notas), start=1):
            print(f"{i} - {nome} | Nota: {nota}")
        print()

def eliminar_aluno():
    if not nomes:
        print("Não existem alunos para eliminar.\n")
    else:
        listar_alunos()
        try:
            indice = int(input("Indique o número do aluno a eliminar: "))
            
            if 1 <= indice <= len(nomes):
                nome_removido = nomes.pop(indice - 1)
                nota_removida = notas.pop(indice - 1)
                print(f"Aluno {nome_removido} eliminado com sucesso!\n")
            else:
                print("Índice inválido.\n")
                
        except ValueError:
            print("Erro! Deve introduzir um número inteiro.\n")

def alterar_aluno():
    if not nomes:
        print("Não existem alunos registados.\n")
    else:
        listar_alunos()
        try:
            indice = int(input("Indique o número do aluno a alterar: "))
            
            if 1 <= indice <= len(nomes):
                novo_nome = input("Introduza o novo nome: ")
                nova_nota = float(input("Introduza a nova nota: "))
                
                if 0 <= nova_nota <= 20:
                    nomes[indice - 1] = novo_nome
                    notas[indice - 1] = nova_nota
                    print("Dados alterados com sucesso!\n")
                else:
                    print("Nota inválida!\n")
            else:
                print("Índice inválido.\n")
                
        except ValueError:
            print("Erro! Dados inválidos.\n")

def estatisticas():
    if not notas:
        print("Não existem dados registados.\n")
    else:
        media = round(np.mean(notas), 2)
        mediana = np.median(notas)
        desvio = round(np.std(notas), 3)
        maior = np.max(notas)
        menor = np.min(notas)
        
        print(f"Média: {media}")
        print(f"Mediana: {mediana}")
        print(f"Desvio padrão: {desvio}")
        print(f"Maior nota: {maior}")
        print(f"Menor nota: {menor}\n")

def exportar_pickle():
    if not nomes:
        print("Não existem dados para guardar.\n")
    else:
        with open("alunos_backup.pkl", "wb") as ficheiro:
            pickle.dump((nomes, notas), ficheiro)
        print("Backup exportado com sucesso!\n")

def importar_pickle():
    global nomes, notas
    try:
        with open("alunos_backup.pkl", "rb") as ficheiro:
            nomes, notas = pickle.load(ficheiro)
        print("Backup importado com sucesso!\n")
    except FileNotFoundError:
        print("Ficheiro não encontrado!\n")

def exportar_word():
    if not nomes:
        print("Não existem dados para exportar.\n")
    else:
        try:
            doc = Document()
            titulo = input("Indique o título do documento: ")
            doc.add_heading(titulo, 0)
            
            for i, (nome, nota) in enumerate(zip(nomes, notas), start=1):
                doc.add_paragraph(f"{i} - {nome} | Nota: {nota}")
            
            doc.save("alunos.docx")
            print("Ficheiro Word gerado com sucesso!\n")
        except Exception as e:
            print(f"Erro ao gerar Word: {e}\n")

def exportar_pdf():
    if not nomes:
        print("Não existem dados para exportar.\n")
    else:
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            titulo = input("Indique o título do documento: ")
            pdf.cell(200, 10, txt=titulo, ln=True, align="C")
            
            for i, (nome, nota) in enumerate(zip(nomes, notas), start=1):
                pdf.cell(200, 10, txt=f"{i} - {nome} | Nota: {nota}", ln=True)
            
            pdf.output("alunos.pdf")
            print("Ficheiro PDF gerado com sucesso!\n")
        except Exception as e:
            print(f"Erro ao gerar PDF: {e}\n")

while True:
    print("Sistema de Gestão de Alunos")
    print("1 - Registar Aluno")
    print("2 - Listar Alunos")
    print("3 - Eliminar Aluno")
    print("4 - Alterar Dados do Aluno")
    print("5 - Ver Estatísticas (NumPy)")
    print("6 - Exportar Backup (Pickle)")
    print("7 - Importar Backup (Pickle)")
    print("8 - Exportar para Word")
    print("9 - Exportar para PDF")
    print("0 - Terminar")

    opcao = input("Escolha uma opção: ")

    if opcao == "1":
        adicionar_aluno()
    elif opcao == "2":
        listar_alunos()
    elif opcao == "3":
        eliminar_aluno()
    elif opcao == "4":
        alterar_aluno()
    elif opcao == "5":
        estatisticas()
    elif opcao == "6":
        exportar_pickle()
    elif opcao == "7":
        importar_pickle()
    elif opcao == "8":
        exportar_word()
    elif opcao == "9":
        exportar_pdf()
    elif opcao == "0":
        print("Programa terminado. Obrigado!")
        break
    else:
        print("Opção inválida!\n")